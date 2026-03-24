"""
Document loader for PolicyLab ingestion pipeline.

Accepts PDF, plain text, Markdown, and DOCX files. Returns structured text
with page/section provenance so every extracted fact can be traced back to
its exact location in the source document.

Provenance is the non-negotiable requirement here. A Brussels staffer asked
"where does this severity score come from?" must be able to trace the answer
to page 47, paragraph 3 of the uploaded regulatory impact assessment. Without
provenance, the pipeline is just another black box.

No external services required — pure Python. PDF extraction uses pypdf (in the
standard pyproject.toml dependencies). DOCX uses python-docx if installed.
"""

from __future__ import annotations

import dataclasses
import os
import re
from pathlib import Path
from typing import Iterator


@dataclasses.dataclass
class TextChunk:
    """A contiguous span of text with its provenance.

    text: the actual text content
    source_id: human-readable reference like "page 3" or "section 2.1"
    char_offset: character offset from the start of the full document text
    chunk_type: "page", "section", "paragraph", or "raw"
    """
    text: str
    source_id: str
    char_offset: int
    chunk_type: str = "raw"

    def __repr__(self) -> str:
        preview = self.text[:60].replace("\n", " ")
        return f"TextChunk(source={self.source_id!r}, text={preview!r}...)"


@dataclasses.dataclass
class LoadedDocument:
    """Result of loading a document.

    full_text: concatenated plain text (use for LLM prompts)
    chunks: list of TextChunk with provenance, ordered by position
    file_path: absolute path to the source file
    file_type: 'pdf', 'txt', 'md', 'docx', or 'unknown'
    n_pages: page count for PDFs; section count for others; None if unknown
    encoding_notes: any warnings from the extraction (e.g. "page 4: garbled OCR")
    """
    full_text: str
    chunks: list[TextChunk]
    file_path: str
    file_type: str
    n_pages: int | None = None
    encoding_notes: list[str] = dataclasses.field(default_factory=list)

    def get_chunk_for_offset(self, char_offset: int) -> TextChunk | None:
        """Return the chunk that contains the given character offset."""
        for chunk in self.chunks:
            end = chunk.char_offset + len(chunk.text)
            if chunk.char_offset <= char_offset < end:
                return chunk
        return None

    def passage_with_context(self, char_offset: int, window: int = 200) -> str:
        """Return text around char_offset with source citation.

        Used when reporting extraction evidence so the analyst can verify it.
        """
        start = max(0, char_offset - window // 2)
        end = min(len(self.full_text), char_offset + window // 2)
        snippet = self.full_text[start:end].strip()
        chunk = self.get_chunk_for_offset(char_offset)
        citation = chunk.source_id if chunk else "unknown location"
        return f"[{citation}] ...{snippet}..."


# ─────────────────────────────────────────────────────────────────────────────
# LOADERS
# ─────────────────────────────────────────────────────────────────────────────

def _clean_text(text: str) -> str:
    """Normalise whitespace and remove control characters."""
    # Collapse multiple blank lines to at most two
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove control chars except newline/tab
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    return text.strip()


def _load_pdf(path: str) -> LoadedDocument:
    """Extract text from a PDF using pypdf, one chunk per page."""
    try:
        import pypdf
    except ImportError:
        raise ImportError(
            "pypdf is required for PDF loading. Install it: pip install pypdf"
        )

    chunks: list[TextChunk] = []
    notes: list[str] = []
    full_parts: list[str] = []
    char_offset = 0

    with open(path, "rb") as f:
        reader = pypdf.PdfReader(f)
        n_pages = len(reader.pages)

        for page_num, page in enumerate(reader.pages, 1):
            try:
                raw = page.extract_text() or ""
                text = _clean_text(raw)
            except Exception as e:
                text = ""
                notes.append(f"page {page_num}: extraction failed ({e})")

            if not text:
                notes.append(f"page {page_num}: no text extracted (may be image-only)")
                continue

            # Add separator between pages
            if full_parts:
                separator = "\n\n"
                char_offset += len(separator)
                full_parts.append(separator)

            chunks.append(TextChunk(
                text=text,
                source_id=f"page {page_num} of {n_pages}",
                char_offset=char_offset,
                chunk_type="page",
            ))
            full_parts.append(text)
            char_offset += len(text)

    full_text = "".join(full_parts)
    return LoadedDocument(
        full_text=full_text,
        chunks=chunks,
        file_path=os.path.abspath(path),
        file_type="pdf",
        n_pages=n_pages,
        encoding_notes=notes,
    )


def _load_text(path: str, file_type: str = "txt") -> LoadedDocument:
    """Load plain text or Markdown, chunked by section headings."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()

    full_text = _clean_text(raw)
    chunks: list[TextChunk] = []

    # Split on Markdown headings (# Title) or blank-line-separated paragraphs
    # Prefer heading-based splits for .md; paragraph-based for .txt
    if file_type == "md":
        # Split on ATX headings: # / ## / ###
        pattern = re.compile(r'^(#{1,3} .+)$', re.MULTILINE)
        positions = [m.start() for m in pattern.finditer(full_text)]
        positions.append(len(full_text))

        if len(positions) <= 1:
            # No headings — fall through to paragraph chunking
            positions = []

        for i, pos in enumerate(positions[:-1]):
            text = full_text[pos:positions[i + 1]].strip()
            if not text:
                continue
            heading_match = pattern.match(text)
            heading = heading_match.group(0) if heading_match else f"section {i+1}"
            chunks.append(TextChunk(
                text=text,
                source_id=heading[:60],
                char_offset=pos,
                chunk_type="section",
            ))

    # If no chunks yet (txt or md with no headings), split by paragraphs
    if not chunks:
        para_re = re.compile(r'\n{2,}')
        cursor = 0
        para_num = 0
        for raw_para in para_re.split(full_text):
            para = raw_para.strip()
            if not para:
                cursor += len(raw_para) + 2  # +2 for the \n\n separator consumed by split
                continue
            para_num += 1
            pos = full_text.find(para, cursor)
            if pos == -1:
                pos = cursor  # fallback: shouldn't happen
            chunks.append(TextChunk(
                text=para,
                source_id=f"paragraph {para_num}",
                char_offset=pos,
                chunk_type="paragraph",
            ))
            cursor = pos + len(para)

    return LoadedDocument(
        full_text=full_text,
        chunks=chunks,
        file_path=os.path.abspath(path),
        file_type=file_type,
        n_pages=len(chunks),
    )


def _load_docx(path: str) -> LoadedDocument:
    """Load DOCX using python-docx, one chunk per paragraph."""
    try:
        import docx  # python-docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX loading. Install it: pip install python-docx"
        )

    doc = docx.Document(path)
    chunks: list[TextChunk] = []
    full_parts: list[str] = []
    char_offset = 0

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        # Headings get a richer source_id
        if para.style.name.startswith("Heading"):
            source_id = f"Heading: {text[:60]}"
            chunk_type = "section"
        else:
            source_id = f"paragraph {i+1}"
            chunk_type = "paragraph"

        if full_parts:
            sep = "\n\n"
            full_parts.append(sep)
            char_offset += len(sep)

        chunks.append(TextChunk(
            text=text,
            source_id=source_id,
            char_offset=char_offset,
            chunk_type=chunk_type,
        ))
        full_parts.append(text)
        char_offset += len(text)

    full_text = "".join(full_parts)
    return LoadedDocument(
        full_text=full_text,
        chunks=chunks,
        file_path=os.path.abspath(path),
        file_type="docx",
        n_pages=len(doc.paragraphs),
    )


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────

def load_document(path: str) -> LoadedDocument:
    """Load a document from disk and return structured text with provenance.

    Supported formats: .pdf, .txt, .md, .markdown, .docx

    Raises FileNotFoundError if the path does not exist.
    Raises ValueError for unsupported formats.
    Raises ImportError if the required library for the format is not installed.
    """
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    elif suffix in (".txt", ""):
        return _load_text(path, "txt")
    elif suffix in (".md", ".markdown"):
        return _load_text(path, "md")
    elif suffix in (".docx", ".doc"):
        return _load_docx(path)
    else:
        raise ValueError(
            f"Unsupported file format: {suffix!r}. "
            f"Supported: .pdf, .txt, .md, .markdown, .docx"
        )


def load_text_string(text: str, name: str = "inline") -> LoadedDocument:
    """Load from a string directly (for testing or inline use).

    Useful when the document is already in memory (e.g. pasted into a UI).
    """
    clean = _clean_text(text)
    chunks = []
    # Track running offset correctly — find(para) from position 0 always
    # returns the FIRST occurrence of the paragraph text, so paragraphs 2, 3...
    # that share any prefix with paragraph 1 get the wrong offset. Use a
    # running cursor that advances past each paragraph instead.
    cursor = 0
    para_num = 0
    for raw_para in re.split(r'\n{2,}', clean):
        para = raw_para.strip()
        if not para:
            # Advance cursor past the blank lines
            cursor += len(raw_para) + 2  # +2 for the \n\n separator
            continue
        para_num += 1
        # Find this paragraph starting AT cursor (not from 0)
        para_start = clean.find(para, cursor)
        if para_start == -1:
            para_start = cursor  # fallback: shouldn't happen after _clean_text
        chunks.append(TextChunk(
            text=para,
            source_id=f"paragraph {para_num}",
            char_offset=para_start,
            chunk_type="paragraph",
        ))
        cursor = para_start + len(para)
    return LoadedDocument(
        full_text=clean,
        chunks=chunks,
        file_path=f"<string:{name}>",
        file_type="txt",
        n_pages=len(chunks),
    )
